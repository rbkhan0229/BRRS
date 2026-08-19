#!/usr/bin/env python3
"""Build the styled DOCX version of the BRRS IEEE working manuscript."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Resolved preset: narrative_proposal. Header pattern: memo_masthead.
PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BODY_FONT = "Calibri"
KOREAN_FONT = "AppleGothic"
ACTIVE_BODY_FONT = BODY_FONT
BODY_SIZE = 11
NAVY = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66727D"
LIGHT_FILL = "F4F6F9"
GRID = "B9C2CA"
CALLOUT_FILL = "EEF3F6"


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=None) -> None:
    font = font or ACTIVE_BODY_FONT
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{script}"), font)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:eastAsia"), "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_bottom_border(paragraph, color=NAVY, size="10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def paragraph_callout(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "160")
    ind.set(qn("w:right"), "120")
    p_pr.append(ind)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_inline_markup(paragraph, text: str, size=None, color=None) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, color=color, bold=True)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), size=size, color=color, font="Consolas")
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size=size, color=color, italic=True)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    def new_num(fmt: str, text: str) -> int:
        abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
        abstract_id = max(abstract_ids, default=-1) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "279")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, jc, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
        num_id = max(num_ids, default=0) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return new_num("bullet", "•"), new_num("decimal", "%1.")


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_styles(doc: Document, language: str = "en") -> None:
    font_name = KOREAN_FONT if language == "ko" else BODY_FONT
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{script}"), font_name)
    normal.font.size = Pt(BODY_SIZE)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = font_name
        for script in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{script}"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = font_name
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        caption._element.rPr.rFonts.set(qn(f"w:{script}"), font_name)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False


def configure_page(doc: Document, language: str = "en") -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("BRRS | 한글 검토용 원고" if language == "ko" else "BRRS | Working manuscript")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    r = p.add_run("IEEE IoT Journal 초안 v0.1" if language == "ko" else "IEEE IoT Journal draft v0.1")
    set_run_font(r, size=8.5, color=MUTED)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5))
    paragraph_bottom_border(p, color="B8C3CB", size="4")

    first = section.first_page_header
    p = first.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "IEEE INTERNET OF THINGS JOURNAL | 한글 검토용 초안"
        if language == "ko"
        else "IEEE INTERNET OF THINGS JOURNAL | WORKING DRAFT"
    )
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    paragraph_bottom_border(p, color=NAVY, size="8")

    for footer in (section.footer, section.first_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(
            "초안 v0.1  |  2026-08-13  |  페이지 "
            if language == "ko"
            else "Draft v0.1  |  13 Aug 2026  |  Page "
        )
        set_run_font(r, size=8.5, color=MUTED)
        add_field(p, "PAGE")


def add_title_block(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=20, color=NAVY, bold=True)


def table_widths(table_no: int, cols: int) -> list[int]:
    patterns = {
        1: [1450, 2150, 2600, 3160],
        2: [1000, 2500, 1450, 1800, 2610],
        3: [900, 1700, 1200, 1500, 4060],
        4: [850, 1650, 1450, 1500, 3910],
        5: [650, 950, 1700, 1100, 1850, 3110],
        6: [2200, 3000, 4160],
    }
    widths = patterns.get(table_no)
    if widths and len(widths) == cols:
        return widths
    base = CONTENT_WIDTH_DXA // cols
    return [base] * (cols - 1) + [CONTENT_WIDTH_DXA - base * (cols - 1)]


def add_markdown_table(doc: Document, lines: list[str], table_no: int) -> None:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [rows[0]] + rows[2:]
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = table_widths(table_no, cols)

    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        set_row_cant_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.text = ""
            if r_idx == 0:
                shade_cell(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            numeric = bool(re.match(r"^[0-9M]|^S[0-9]|^0\.|^1\.", value))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if numeric or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markup(p, value, size=8.6, color=NAVY if r_idx == 0 else None)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True

    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)


def build(md_path: Path, output_path: Path, language: str = "en") -> None:
    global ACTIVE_BODY_FONT
    ACTIVE_BODY_FONT = KOREAN_FONT if language == "ko" else BODY_FONT
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc, language)
    configure_page(doc, language)
    bullet_num, decimal_num = add_numbering_definitions(doc)

    props = doc.core_properties
    props.title = (
        "소형 센서 보고를 위한 비컨 참조형 단축 프리앰블 HRP-UWB TDMA"
        if language == "ko"
        else "Beacon-Referenced Short-Preamble HRP-UWB TDMA for Small Sensor Reports"
    )
    props.subject = (
        "IEEE Internet of Things Journal 한글 검토용 원고"
        if language == "ko"
        else "IEEE Internet of Things Journal working manuscript"
    )
    props.author = "익명 저자" if language == "ko" else "Anonymous Authors"
    props.keywords = "UWB, BRRS, DWM3000, TDMA, preamble"

    i = 0
    table_no = 0
    title_done = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("# ") and not title_done:
            add_title_block(doc, line[2:])
            title_done = True
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:]
            p = doc.add_paragraph(style="Heading 1")
            if text in (
                "V. Results",
                "V. 결과",
                "Appendix A. Draft Claim Ledger",
                "부록 A. 초안 주장-근거 대응표",
            ):
                p.paragraph_format.page_break_before = True
            add_inline_markup(p, text)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markup(p, line[4:])
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.2
            paragraph_callout(p)
            add_inline_markup(p, line[2:], color=NAVY)
            i += 1
            continue

        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                caption, rel_path = m.groups()
                image_path = (md_path.parent / rel_path).resolve()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(0)
                p.add_run().add_picture(str(image_path), width=Inches(6.45))
                cap = doc.add_paragraph(style="Caption")
                add_inline_markup(cap, caption, size=9, color=MUTED)
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            table_no += 1
            add_markdown_table(doc, table_lines, table_no)
            continue

        if line.startswith("$$") and line.endswith("$$"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.keep_together = True
            set_run_font(p.add_run(line[2:-2]), size=11, color=NAVY, italic=True, font="Cambria Math")
            i += 1
            continue

        if re.match(r"^- ", line):
            p = doc.add_paragraph()
            apply_num(p, bullet_num)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline_markup(p, line[2:])
            i += 1
            continue

        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            apply_num(p, decimal_num)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline_markup(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue

        # First-page metadata and manuscript table captions are named overrides.
        if line.startswith("**Anonymous Authors**") or line.startswith("**익명 저자**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_inline_markup(p, line.replace("  ", ""), size=11, color=MUTED)
            i += 1
            continue
        if line.startswith("**Working manuscript") or line.startswith("**IEEE Internet of Things Journal 투고용"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            add_inline_markup(p, line, size=9.5, color=MUTED)
            paragraph_bottom_border(p, color=NAVY, size="8")
            i += 1
            continue
        if line.startswith("**TABLE ") or line.startswith("**표 "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            add_inline_markup(p, line, size=8.7, color=NAVY)
            i += 1
            continue

        p = doc.add_paragraph()
        if line.startswith("**Abstract—**") or line.startswith("**초록—**"):
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(3)
        elif line.startswith("**Index Terms—**") or line.startswith("**색인어—**"):
            p.paragraph_format.space_after = Pt(10)
        elif line.startswith("[") and re.match(r"^\[\d+\]", line):
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.28)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.08
        else:
            p.paragraph_format.first_line_indent = Inches(0.2)
        add_inline_markup(p, line)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--language", choices=("en", "ko"), default="en")
    args = parser.parse_args()
    build(args.markdown, args.output, args.language)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
