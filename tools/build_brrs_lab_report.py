from __future__ import annotations

import csv
import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/songchieon/Desktop/DWM3000/logs/result1")
OUT_DIR = Path("/Users/songchieon/Desktop/DWM3000/DW3_QM33_SDK_1.0.2/reports")
ASSET_DIR = OUT_DIR / "assets_brrs_lab_report"
DOCX_PATH = OUT_DIR / "brrs_lab_report_20260702.docx"

FONT_KR = "Apple SD Gothic Neo"
FONT_LATIN = "Arial"
PIL_FONT_REG = "/System/Library/Fonts/AppleSDGothicNeo.ttc"

BLUE = "1F4E79"
DARK = "1F2937"
MUTED = "5B6770"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
PALE_GREEN = "E2F0D9"
WHITE = "FFFFFF"
GRID = "D6DCE2"


@dataclass
class Exp1Row:
    plen: int
    rx: int
    expected: int
    miss: int
    per: float
    err: int
    timeout: int
    rxerr: int
    latency_avg: int
    uwb_offset: int


@dataclass
class MarginRow:
    condition: str
    rx: int
    expected: int
    miss: int
    per: float
    timeout: int
    rxerr: int


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def parse_exp1(md_path: Path) -> list[Exp1Row]:
    text = md_path.read_text(encoding="utf-8")
    rows: list[Exp1Row] = []
    sections = re.split(r"\n---\n", text)
    for section in sections:
        h = re.search(r"##\s+(\d+)sym", section)
        if not h:
            continue
        plen = int(h.group(1))
        m = re.search(
            r"N2:\s+rx=(\d+)\s+expected=(\d+)\s+miss=(\d+)\s+PER=([\d.]+)%\s+err=(\d+)",
            section,
        )
        t = re.search(r"RX timeouts=(\d+)\s+RX errors=(\d+)", section)
        lat = re.search(r"N2:\s+min=\d+us\s+max=\d+us\s+avg=(\d+)us\s+\(n=\d+\)", section)
        offsets = re.findall(r"N2:\s+min=\d+us\s+max=\d+us\s+avg=(\d+)us\s+\(n=\d+\)", section)
        if m and t:
            latency_avg = int(lat.group(1)) if lat else 0
            uwb_offset = int(offsets[-1]) if len(offsets) >= 3 else 0
            rows.append(
                Exp1Row(
                    plen=plen,
                    rx=int(m.group(1)),
                    expected=int(m.group(2)),
                    miss=int(m.group(3)),
                    per=float(m.group(4)),
                    err=int(m.group(5)),
                    timeout=int(t.group(1)),
                    rxerr=int(t.group(2)),
                    latency_avg=latency_avg,
                    uwb_offset=uwb_offset,
                )
            )
    return sorted(rows, key=lambda r: r.plen)


def parse_margin(md_path: Path) -> list[MarginRow]:
    text = md_path.read_text(encoding="utf-8")
    rows: list[MarginRow] = []
    sections = re.split(r"\n---\n", text)
    for section in sections:
        h = re.search(r"##\s+32sym\s+-\s+(.+)", section)
        if not h:
            continue
        condition = h.group(1).strip()
        m = re.search(
            r"N2:\s+rx=(\d+)\s+expected=(\d+)\s+miss=(\d+)\s+PER=([\d.]+)%\s+err=(\d+)",
            section,
        )
        t = re.search(r"RX timeouts=(\d+)\s+RX errors=(\d+)", section)
        if m and t:
            label = "Lead margin only" if "lead" in condition.lower() else "Tail margin only"
            rows.append(
                MarginRow(
                    condition=label,
                    rx=int(m.group(1)),
                    expected=int(m.group(2)),
                    miss=int(m.group(3)),
                    per=float(m.group(4)),
                    timeout=int(t.group(1)),
                    rxerr=int(t.group(2)),
                )
            )
    return rows


def parse_exp2_summary(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return sorted(list(csv.DictReader(f)), key=lambda row: int(row["plen"]))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    try:
        return ImageFont.truetype(PIL_FONT_REG, size=size, index=8 if bold else 0)
    except Exception:
        return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt, fill):
    bbox = draw.textbbox((0, 0), text, font=fnt)
    draw.text((xy[0] - (bbox[2] - bbox[0]) / 2, xy[1] - (bbox[3] - bbox[1]) / 2), text, font=fnt, fill=fill)


def make_bar_chart_exp1(rows: list[Exp1Row], out: Path) -> None:
    w, h = 1400, 760
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(30, True)
    small_f = font(24)
    tiny_f = font(20)

    d.text((70, 44), "Experiment 1: Preamble Length vs. PER", font=title_f, fill="#111827")
    d.text((72, 96), "1000 cycles, Node N2, delayed-RX/TX, office environment", font=small_f, fill="#4B5563")
    left, top, right, bottom = 120, 170, 1320, 610
    d.line((left, bottom, right, bottom), fill="#111827", width=3)
    d.line((left, top, left, bottom), fill="#111827", width=3)
    max_y = 1.0
    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = bottom - (tick / max_y) * (bottom - top)
        d.line((left, y, right, y), fill="#E5E7EB", width=2)
        d.text((50, y - 14), f"{tick:.2f}%", font=tiny_f, fill="#374151")
    colors = ["#D95F02" if r.per > 0 else "#2E7D32" for r in rows]
    gap = (right - left) / len(rows)
    for i, row in enumerate(rows):
        cx = left + gap * i + gap / 2
        bar_w = gap * 0.42
        value = max(row.per, 0.02 if row.per == 0 else row.per)
        y = bottom - (value / max_y) * (bottom - top)
        d.rounded_rectangle((cx - bar_w / 2, y, cx + bar_w / 2, bottom), radius=10, fill=colors[i])
        d.text((cx - 42, bottom + 24), f"{row.plen}", font=label_f, fill="#111827")
        d.text((cx - 50, bottom + 66), "sym", font=tiny_f, fill="#6B7280")
        d.text((cx - 58, y - 38), f"{row.per:.2f}%", font=label_f, fill="#111827")
        d.text((cx - 65, y - 66), f"miss {row.miss}", font=tiny_f, fill="#6B7280")
    d.text((w / 2 - 145, h - 58), "Preamble length (symbols)", font=small_f, fill="#374151")
    d.text((10, 362), "PER", font=small_f, fill="#374151")
    img.save(out)


def make_margin_chart(rows: list[MarginRow], out: Path) -> None:
    w, h = 1400, 760
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(30, True)
    small_f = font(24)
    tiny_f = font(20)
    d.text((70, 44), "Lead Margin vs. Tail Margin (32 symbols)", font=title_f, fill="#111827")
    d.text((72, 96), "PER comparison: opening earlier vs. closing later", font=small_f, fill="#4B5563")
    left, top, right, bottom = 160, 170, 1280, 610
    d.line((left, bottom, right, bottom), fill="#111827", width=3)
    d.line((left, top, left, bottom), fill="#111827", width=3)
    max_y = 100.0
    for tick in [0, 25, 50, 75, 100]:
        y = bottom - (tick / max_y) * (bottom - top)
        d.line((left, y, right, y), fill="#E5E7EB", width=2)
        d.text((70, y - 14), f"{tick}%", font=tiny_f, fill="#374151")
    gap = (right - left) / len(rows)
    for i, row in enumerate(rows):
        cx = left + gap * i + gap / 2
        bar_w = 230
        color = "#2E7D32" if row.per < 1 else "#C00000"
        value = max(row.per, 0.7)
        y = bottom - (value / max_y) * (bottom - top)
        d.rounded_rectangle((cx - bar_w / 2, y, cx + bar_w / 2, bottom), radius=12, fill=color)
        label = "Lead only" if "Lead" in row.condition else "Tail only"
        draw_center(d, (int(cx), bottom + 38), label, label_f, "#111827")
        draw_center(d, (int(cx), bottom + 78), f"rx {row.rx}/{row.expected}", small_f, "#4B5563")
        draw_center(d, (int(cx), int(y - 42)), f"{row.per:.2f}%", label_f, "#111827")
        draw_center(d, (int(cx), int(y - 74)), f"timeouts {row.timeout}", tiny_f, "#6B7280")
    d.text((18, 360), "PER", font=small_f, fill="#374151")
    img.save(out)


def make_exp2_chart(rows: list[dict[str, str]], out: Path) -> None:
    w, h = 1500, 820
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(26, True)
    small_f = font(23)
    tiny_f = font(19)

    xvals = [int(r["plen"]) for r in rows]
    measured = [float(r["fp_snr_db_mean"]) for r in rows]
    model = [float(r["processing_gain_db"]) for r in rows]
    ymin = math.floor(min(measured + model) - 2)
    ymax = math.ceil(max(measured + model) + 1)
    left, top, right, bottom = 140, 170, 1360, 640
    d.text((70, 44), "Experiment 2: First-path SNR vs. Preamble Length", font=title_f, fill="#111827")
    d.text((72, 96), "Measured raw CIR FP-SNR compared with processing-gain model Gp(M)", font=small_f, fill="#4B5563")
    d.line((left, bottom, right, bottom), fill="#111827", width=3)
    d.line((left, top, left, bottom), fill="#111827", width=3)
    for tick in range(ymin, ymax + 1, 4):
        y = bottom - ((tick - ymin) / (ymax - ymin)) * (bottom - top)
        d.line((left, y, right, y), fill="#E5E7EB", width=2)
        d.text((60, y - 13), f"{tick}", font=tiny_f, fill="#374151")

    def xcoord(x: int) -> float:
        return left + (math.log2(x) - math.log2(min(xvals))) / (math.log2(max(xvals)) - math.log2(min(xvals))) * (right - left)

    def ycoord(yv: float) -> float:
        return bottom - ((yv - ymin) / (ymax - ymin)) * (bottom - top)

    measured_pts = [(xcoord(x), ycoord(y)) for x, y in zip(xvals, measured)]
    model_pts = [(xcoord(x), ycoord(y)) for x, y in zip(xvals, model)]
    d.line(measured_pts, fill="#1F77B4", width=5)
    d.line(model_pts, fill="#C00000", width=4)
    for x, y, m, row in zip(xvals, measured, measured_pts, rows):
        d.ellipse((m[0] - 9, m[1] - 9, m[0] + 9, m[1] + 9), fill="#1F77B4")
        d.text((m[0] - 42, m[1] - 42), f"{y:.1f}", font=label_f, fill="#111827")
        status = row["status"]
        d.text((m[0] - 34, m[1] + 16), status, font=tiny_f, fill="#C00000" if status == "FAIL" else "#2E7D32")
    for p in model_pts:
        d.ellipse((p[0] - 7, p[1] - 7, p[0] + 7, p[1] + 7), fill="#C00000")
    for x in xvals:
        xx = xcoord(x)
        d.line((xx, bottom, xx, bottom + 8), fill="#111827", width=2)
        draw_center(d, (int(xx), bottom + 38), str(x), label_f, "#111827")
    d.text((w / 2 - 145, h - 70), "Preamble length M (symbols)", font=small_f, fill="#374151")
    d.text((18, 380), "FP-SNR (dB)", font=small_f, fill="#374151")
    d.rectangle((1010, 175, 1040, 195), fill="#1F77B4")
    d.text((1052, 170), "Measured FP-SNR", font=small_f, fill="#111827")
    d.rectangle((1010, 213, 1040, 233), fill="#C00000")
    d.text((1052, 208), "Gp model", font=small_f, fill="#111827")
    img.save(out)


def set_run_font(run, name=FONT_KR, size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style_font(style, size, color="000000", bold=False):
    style.font.name = FONT_KR
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_KR)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_KR)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KR)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5) -> None:
    cell.text = ""
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, color="111827", bold=bold)


def set_table_borders(table, color=GRID):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_width(cell, width_in: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.bold = True
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=MUTED)


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=12, color=DARK, bold=True)
        rest = text[len(bold_prefix) :]
        r = p.add_run(rest)
        set_run_font(r, size=12, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12, color=DARK)
    return p


def add_callout(doc: Document, title: str, bullets: list[str], fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="B7C9D8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=13, color=BLUE, bold=True)
    for bullet in bullets:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(bullet)
        set_run_font(r, size=11.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.65, 4.85)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)
    for i, (k, v) in enumerate(rows):
        set_cell_width(table.cell(i, 0), widths[0])
        set_cell_width(table.cell(i, 1), widths[1])
        set_cell_text(table.cell(i, 0), k, bold=True, fill=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.8)
        set_cell_text(table.cell(i, 1), v, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for j, h in enumerate(headers):
        set_cell_width(table.cell(0, j), widths[j])
        set_cell_text(table.cell(0, j), h, bold=True, fill=LIGHT_BLUE, size=10.5)
    for row in rows:
        cells = table.add_row().cells
        for j, v in enumerate(row):
            set_cell_width(cells[j], widths[j])
            fill = None
            if "FAIL" in v or "100.00%" in v:
                fill = PALE_RED
            elif "PASS" in v or v == "0.00%":
                fill = PALE_GREEN
            set_cell_text(cells[j], v, fill=fill, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_compact_definition_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table)
    set_cell_width(table.cell(0, 0), 1.35)
    set_cell_width(table.cell(0, 1), 5.15)
    set_cell_text(table.cell(0, 0), "지표", bold=True, fill=LIGHT_GRAY, size=9.5)
    set_cell_text(table.cell(0, 1), "구성", bold=True, fill=LIGHT_GRAY, size=9.5)
    for label, desc in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 1.35)
        set_cell_width(cells[1], 5.15)
        set_cell_text(cells[0], label, bold=True, fill=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.2)
        set_cell_text(cells[1], desc, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    set_style_font(doc.styles["Normal"], 12, DARK)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    set_style_font(doc.styles["Heading 1"], 18, BLUE, True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(doc.styles["Heading 2"], 14, BLUE, True)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(doc.styles["List Bullet"], 11.5, DARK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("BRRS Lab Meeting Report")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DWM3000 BRRS experiment summary")
    set_run_font(r, size=9, color=MUTED)


def add_page_break(doc: Document):
    doc.add_page_break()


def build_docx(exp1: list[Exp1Row], margin_rows: list[MarginRow], exp2: list[dict[str, str]], charts: dict[str, Path]):
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("BRRS 프리앰블 축소 실험 결과 보고")
    set_run_font(r, size=24, color="000000", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("랩미팅 보고용 | DWM3000 delayed-RX/TX 기반 | 2026-07-02")
    set_run_font(r, size=12.5, color=MUTED)

    add_callout(
        doc,
        "핵심 결론",
        [
            "Exp1: 사무실 1 m 조건에서 64/128/256 sym은 PER 0.00%, 32 sym은 PER 0.20%로 1% 기준을 만족했다.",
            "Margin 비교: 32 sym에서는 tail margin보다 lead margin이 결정적이었다. Tail-only는 RX timeout 1000회로 수신 실패했다.",
            "Exp2: first-path SNR은 프리앰블 길이에 따라 증가했고, 64 sym 이상에서는 doubling마다 약 3 dB 증가해 처리 이득 모델의 기울기와 대체로 일치했다.",
            "해석: 짧은 프리앰블의 채널 airtime 감소 효과는 유지되지만, 안정 수신을 위해 RX window를 충분히 일찍 여는 lead margin 확보가 필요하다.",
        ],
    )

    add_heading(doc, "1. 실험 환경", 1)
    add_kv_table(
        doc,
        [
            ("장소/거리", "사무실 환경, 노드 간 약 1 m"),
            ("장비/노드", "DWM3000 기반 INIT 1대 + Normal Node 2 1대"),
            ("공통 설정", "PSDU 127 bytes, delayed-RX/TX, 1000 cycles"),
            ("프리앰블", "32, 64, 128, 256 symbols"),
            ("최종 margin", "lead margin 100 us, tail margin 100 us 기준으로 Exp1 최종 결과 정리"),
        ],
    )
    env_img = ROOT / "exp_office_environment.jpeg"
    if env_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(env_img), width=Inches(3.15))
        add_caption(doc, "그림 1. 사무실 실험 환경 사진")

    add_page_break(doc)
    add_heading(doc, "2. Lead Margin / Tail Margin 영향 비교", 1)
    add_body(
        doc,
        "관찰 요약: 32 sym처럼 프리앰블이 짧은 조건에서는 수신기를 늦게 닫는 것보다 예정 RMARKER보다 충분히 일찍 켜는 것이 더 중요했다.",
        bold_prefix="관찰 요약:",
    )
    add_data_table(
        doc,
        ["조건", "RX/Expected", "Miss", "PER", "Timeout", "RX Error", "판정"],
        [
            [
                r.condition.replace("Lead margin only", "Lead only").replace("Tail margin only", "Tail only"),
                f"{r.rx}/{r.expected}",
                str(r.miss),
                f"{r.per:.2f}%",
                str(r.timeout),
                str(r.rxerr),
                "PASS" if r.per <= 1.0 else "FAIL",
            ]
            for r in margin_rows
        ],
        [1.25, 1.25, 0.7, 0.8, 0.85, 0.85, 0.8],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["margin"]), width=Inches(6.8))
    add_caption(doc, "그림 2. 32 sym 조건에서 lead-only와 tail-only의 PER 비교")
    add_callout(
        doc,
        "해석",
        [
            "Lead margin은 RX_ON 시점을 앞당겨 preamble/SFD를 놓치지 않게 한다.",
            "Tail margin은 이미 검출이 시작된 프레임의 뒤쪽 여유를 주는 역할이므로, RX_ON이 늦어 preamble을 놓친 경우에는 회복 효과가 제한적이다.",
            "이번 결과에서 tail-only가 100% timeout인 것은 수신 종료 문제가 아니라 수신 시작 시점 문제가 지배적이었음을 의미한다.",
        ],
        fill=PALE_YELLOW,
    )

    add_page_break(doc)
    add_heading(doc, "3. Experiment 1 최종 결과: 프리앰블 길이별 PER", 1)
    add_body(
        doc,
        "목적: delayed-RX 기반에서 프리앰블을 줄였을 때 PER이 1% 이하로 유지되는 최소 길이를 확인한다.",
        bold_prefix="목적:",
    )
    add_data_table(
        doc,
        ["Preamble", "RX/Expected", "Miss", "PER", "Timeout", "Latency avg", "SYNC-DATA offset"],
        [
            [
                f"{r.plen} sym",
                f"{r.rx}/{r.expected}",
                str(r.miss),
                f"{r.per:.2f}%",
                str(r.timeout),
                f"{r.latency_avg} us",
                f"{r.uwb_offset} us",
            ]
            for r in exp1
        ],
        [1.0, 1.2, 0.65, 0.8, 0.85, 1.05, 1.05],
    )
    add_compact_definition_table(
        doc,
        [
            (
                "Latency avg",
                "rx_ts - tx_ts. Normal의 DATA 송신 준비 시각부터 INIT의 DATA 수신 처리 시각까지의 CPU timer 기반 차이이며, 순수 전파 지연은 아니다.",
            ),
            (
                "SYNC-DATA offset",
                "INIT의 SYNC TX UWB timestamp부터 N2 DATA RX UWB timestamp까지의 차이. 실제 슬롯 도착 시각 확인용 UWB timestamp 값이다.",
            ),
        ],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["exp1"]), width=Inches(6.9))
    add_caption(doc, "그림 3. 프리앰블 길이별 PER 최종 결과")
    add_callout(
        doc,
        "Exp1 결론",
        [
            "64 sym 이상은 이번 사무실 1 m 조건에서 1000/1000 수신으로 안정적이었다.",
            "32 sym도 PER 0.20%로 목표 기준(PER <= 1%)은 만족하지만, miss가 발생하므로 위치/차폐/높이 변화에 민감한 후보로 보는 것이 타당하다.",
            "프리앰블이 길어질수록 SYNC-DATA offset과 latency가 증가하므로, 안정성과 airtime 절감 사이의 균형점은 32-64 sym 구간에 있다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "4. Experiment 2 결과: 프리앰블 길이별 CIR 품질", 1)
    add_body(
        doc,
        "목적: 각 프리앰블 길이에서 raw CIR 기반 first-path SNR을 산출하고, 처리 이득 모델 Gp(M)=10log10(127M)의 경향과 비교한다.",
        bold_prefix="목적:",
    )
    add_data_table(
        doc,
        ["Preamble", "n", "Status", "FP-SNR mean", "FP-SNR median", "Model Gp", "비고"],
        [
            [
                f"{int(r['plen'])} sym",
                f"{int(r['n'])}/{int(r['expected_n'])}",
                r["status"],
                f"{float(r['fp_snr_db_mean']):.2f} dB",
                f"{float(r['fp_snr_db_median']):.2f} dB",
                f"{float(r['processing_gain_db']):.2f} dB",
                "1 sample missing" if r["status"] == "FAIL" else "-",
            ]
            for r in exp2
        ],
        [0.95, 0.75, 0.75, 1.1, 1.1, 0.95, 1.3],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["exp2"]), width=Inches(7.0))
    add_caption(doc, "그림 4. 프리앰블 길이별 first-path SNR과 처리 이득 모델 비교")
    add_callout(
        doc,
        "Exp2 해석",
        [
            "64 -> 128 sym은 +3.45 dB, 128 -> 256 sym은 +3.09 dB로, doubling당 약 +3 dB인 모델 기울기와 잘 맞는다.",
            "32 -> 64 sym 증가는 +7.63 dB로 더 크다. 이는 32 sym이 검출 한계에 가까워 환경/타이밍 영향이 크게 반영된 것으로 해석된다.",
        ],
        fill=PALE_YELLOW,
    )

    add_heading(doc, "5. 종합 결론 및 다음 실험 제안", 1)
    add_callout(
        doc,
        "보고 결론",
        [
            "Delayed-RX 구조에서는 프리앰블 전체를 길게 유지하는 대신, RX_ON 시점을 충분히 앞당기는 lead margin이 짧은 프리앰블 안정성에 직접적으로 기여했다.",
            "현재 사무실 1 m 조건에서는 64 sym이 안정성과 airtime 절감의 보수적 선택이고, 32 sym은 목표 PER은 만족하지만 반복/환경 검증이 필요한 공격적 선택이다.",
            "CIR 결과는 64 sym 이상에서 처리 이득 모델의 증가 추세를 뒷받침한다. 따라서 Exp1의 PER 결과를 Exp2의 FP-SNR 변화로 설명할 수 있다.",
        ],
    )
    add_heading(doc, "다음 단계", 2)
    next_steps = [
        "32/64 sym을 중심으로 거리(0.5, 1, 2, 3 m)와 높이/차폐 조건을 나누어 반복 측정한다.",
        "차량 범퍼-실내 노드 시나리오에서 32 sym의 PER과 FP-SNR이 유지되는지 확인한다.",
        "Lead margin을 0/50/100/200 us로 sweep하여 RX window 증가량과 PER 개선량의 trade-off를 정량화한다.",
    ]
    for step in next_steps:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(step)
        set_run_font(r, size=12, color=DARK)

    add_heading(doc, "부록: 원자료 위치", 2)
    add_kv_table(
        doc,
        [
            ("Exp1 원문", str(ROOT / "exp1" / "exp1_result.md")),
            ("Margin 비교 원문", str(ROOT / "exp1" / "lead vs tail.md")),
            ("Exp2 summary", str(ROOT / "exp2" / "exp2_cir_office_1m_summary.csv")),
            ("실험환경 사진", str(ROOT / "exp_office_environment.jpeg")),
        ],
        widths=(1.55, 4.95),
    )

    doc.save(DOCX_PATH)


def main():
    ensure_dirs()
    exp1 = parse_exp1(ROOT / "exp1" / "exp1_result.md")
    margin_rows = parse_margin(ROOT / "exp1" / "lead vs tail.md")
    exp2 = parse_exp2_summary(ROOT / "exp2" / "exp2_cir_office_1m_summary.csv")

    charts = {
        "exp1": ASSET_DIR / "exp1_per_chart.png",
        "margin": ASSET_DIR / "lead_tail_margin_chart.png",
        "exp2": ASSET_DIR / "exp2_fp_snr_chart.png",
    }
    make_bar_chart_exp1(exp1, charts["exp1"])
    make_margin_chart(margin_rows, charts["margin"])
    make_exp2_chart(exp2, charts["exp2"])
    build_docx(exp1, margin_rows, exp2, charts)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
